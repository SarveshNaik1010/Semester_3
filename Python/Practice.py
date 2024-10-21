from docx import Document

# Create a new Word document
doc = Document()

# Set styles for the document
doc.styles['Normal'].font.name = 'Times New Roman'
doc.styles['Normal'].font.size = 12

# Title and Formatting
doc.add_heading('McDonald\'s Customer Sentiment Analysis', level=1)

# Project Objective
doc.add_heading('Project Objective', level=2)
doc.add_paragraph(
    "The objective of this project is to analyze customer reviews "
    "to classify sentiments as Positive, Negative, or Neutral. "
    "The results help in understanding customer satisfaction levels, "
    "identifying potential service issues, and leveraging positive feedback for improvement."
)

# Project Description
doc.add_heading('Project Details/Project Description', level=2)

doc.add_heading('1. Dataset Description', level=3)
doc.add_paragraph(
    "The dataset used in this project consists of customer reviews collected from "
    "various McDonald’s stores. The key features include review text, star rating, "
    "location, and review time. The dataset aims to capture customer opinions, complaints, "
    "or compliments towards McDonald's services."
)

doc.add_heading('2. Relative Theory', level=3)
doc.add_paragraph(
    "Sentiment analysis is an NLP task that involves classifying emotions expressed in text. "
    "It uses techniques like tokenization, lemmatization, and machine learning models to identify "
    "whether the sentiment is positive, negative, or neutral. Businesses, like McDonald's, use "
    "this approach to evaluate customer feedback and improve their services."
)

doc.add_heading('3. Implementation and Methodology', level=3)
doc.add_paragraph(
    "1. Data Preprocessing: Reviews were cleaned by lowercasing, tokenization, stopword removal, "
    "and lemmatization using spaCy.\n"
    "2. Sentiment Classification: A sentiment analysis model (VADER) was used to classify reviews "
    "into Positive, Negative, and Neutral sentiments based on polarity scores.\n"
    "3. Visualization: Seaborn and Matplotlib libraries were used to visualize the sentiment distribution.\n"
    "4. Results: A bar graph was plotted to show the distribution of different sentiments."
)

# Project Outcome
doc.add_heading('Project Outcome', level=2)
doc.add_paragraph(
    "The sentiment analysis provided valuable insights into customer experiences at McDonald's. "
    "It helped identify recurring issues like delayed service and highlighted positive aspects such as friendly staff. "
    "This analysis allows McDonald's to take targeted actions for customer satisfaction and operational improvements."
)

# Conclusion
doc.add_heading('Conclusion', level=2)
doc.add_paragraph(
    "This project provided hands-on experience with natural language processing and sentiment analysis. "
    "From the dataset, we learned how to handle textual data, clean it, and apply machine learning techniques "
    "to classify sentiments. The analysis highlighted key customer feedback patterns, enabling businesses to respond "
    "more effectively to customer needs."
)

# Save the document
file_path = "file.docx"
doc.save(file_path)

file_path